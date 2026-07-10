"""Generate .docx translation documents."""

from io import BytesIO

from docx import Document
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.shared import Pt

from app.constants.languages import get_language


def _set_rtl(paragraph) -> None:
    """为段落设置 RTL 双向方向（<w:bidi/>）。"""
    pPr = paragraph._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:bidi {nsdecls("w")}/>'))


def generate_translation_docx(
    source_text: str,
    translated_text: str,
    risk_annotations: list[dict],
    language: str,
) -> bytes:
    """Generate a standard-format .docx with source text and translation.

    Returns the .docx content as bytes.
    """
    if not translated_text:
        raise ValueError("translated_text is required")

    doc = Document()

    # Set default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Times New Roman"
    font.size = Pt(12)

    # --- Source text section ---
    doc.add_heading("【原文】", level=2)
    para = doc.add_paragraph(source_text)
    para.style = doc.styles["Normal"]

    # --- Separator ---
    doc.add_paragraph("─" * 40)

    # --- Translation section ---
    doc.add_heading("【译文】", level=2)
    para_tr = doc.add_paragraph(translated_text)
    para_tr.style = doc.styles["Normal"]

    # RTL 语言（阿拉伯语/乌尔都语）译文段落设双向方向
    lang_info = get_language(language)
    if lang_info is not None and lang_info.direction == "rtl":
        _set_rtl(para_tr)

    # Risk annotations are no longer included as Word comments

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
