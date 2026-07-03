"""Tests for .docx export service."""

import io
from io import BytesIO

import pytest
from docx import Document as DocxDocument
from docx import Document
from docx.oxml.ns import qn


@pytest.fixture
def sample_source() -> str:
    return "坚持以人民为中心的发展思想。"


@pytest.fixture
def sample_translation() -> str:
    return "Adhere to the people-centered development philosophy."


@pytest.fixture
def sample_annotations() -> list[dict]:
    return [
        {
            "phrase": "people-centered",
            "risk_level": "medium",
            "risk_type": "political_sensitivity",
            "explanation": "部分西方媒体将其与民粹主义关联",
        }
    ]


class TestGenerateTranslationDocx:
    def test_returns_non_empty_bytes(self, sample_source, sample_translation, sample_annotations):
        """Verify the function returns a non-empty bytes object."""
        from app.services.export_docx import generate_translation_docx

        result = generate_translation_docx(
            source_text=sample_source,
            translated_text=sample_translation,
            risk_annotations=sample_annotations,
            language="en-GB",
        )
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_docx_contains_source_and_translation(self, sample_source, sample_translation, sample_annotations):
        """Verify the generated .docx contains both source and translation paragraphs."""
        from app.services.export_docx import generate_translation_docx

        result = generate_translation_docx(
            source_text=sample_source,
            translated_text=sample_translation,
            risk_annotations=sample_annotations,
            language="en-GB",
        )
        doc = DocxDocument(io.BytesIO(result))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "坚持以人民为中心的发展思想" in text
        assert "people-centered development" in text

    def test_docx_has_headings(self, sample_source, sample_translation, sample_annotations):
        """Verify the document has at least one heading (原文/译文)."""
        from app.services.export_docx import generate_translation_docx

        result = generate_translation_docx(
            source_text=sample_source,
            translated_text=sample_translation,
            risk_annotations=sample_annotations,
            language="en-GB",
        )
        doc = DocxDocument(io.BytesIO(result))
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) >= 2

    def test_annotations_become_comments(self, sample_source, sample_translation, sample_annotations):
        """Verify risk annotations are added as Word comments on matching text."""
        from app.services.export_docx import generate_translation_docx

        result = generate_translation_docx(
            source_text=sample_source,
            translated_text=sample_translation,
            risk_annotations=sample_annotations,
            language="en-GB",
        )
        doc = DocxDocument(io.BytesIO(result))

        comment_count = 0
        for rel in doc.part.rels.values():
            if "comment" in str(rel.reltype).lower():
                comment_count += 1
        # At least one comment relationship exists
        assert comment_count >= 1

    def test_empty_annotations_list(self, sample_source, sample_translation):
        """Verify the function handles an empty annotations list."""
        from app.services.export_docx import generate_translation_docx

        result = generate_translation_docx(
            source_text=sample_source,
            translated_text=sample_translation,
            risk_annotations=[],
            language="en-GB",
        )
        doc = DocxDocument(io.BytesIO(result))
        text = "\n".join(p.text for p in doc.paragraphs)
        assert "原文" in text
        assert "译文" in text

    def test_empty_translation_raises(self, sample_source, sample_annotations):
        """Verify empty translation raises ValueError."""
        from app.services.export_docx import generate_translation_docx

        with pytest.raises(ValueError, match="translated_text is required"):
            generate_translation_docx(
                source_text=sample_source,
                translated_text="",
                risk_annotations=sample_annotations,
                language="en-GB",
            )


@pytest.fixture
async def async_client():
    """Create a test client for the FastAPI app."""
    from app.main import app
    from httpx import AsyncClient, ASGITransport

    transport = ASGITransport(app=app)
    async with AsyncClient(transport=transport, base_url="http://test") as client:
        yield client


class TestExportApi:
    """Integration tests for POST /api/export/docx."""

    async def test_export_returns_docx(self, async_client, sample_source, sample_translation, sample_annotations):
        """Verify the endpoint returns a .docx file."""
        resp = await async_client.post(
            "/api/export/docx",
            json={
                "source_text": sample_source,
                "translated_text": sample_translation,
                "risk_annotations": sample_annotations,
                "language": "en-GB",
            },
        )
        assert resp.status_code == 200
        assert resp.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert len(resp.content) > 0

    async def test_export_empty_translation_returns_400(self, async_client, sample_source):
        """Verify the endpoint rejects empty translation."""
        resp = await async_client.post(
            "/api/export/docx",
            json={
                "source_text": sample_source,
                "translated_text": "",
                "risk_annotations": [],
                "language": "en-GB",
            },
        )
        assert resp.status_code == 400

    async def test_export_no_annotations(self, async_client, sample_source, sample_translation):
        """Verify the endpoint works with no risk annotations."""
        resp = await async_client.post(
            "/api/export/docx",
            json={
                "source_text": sample_source,
                "translated_text": sample_translation,
                "risk_annotations": [],
                "language": "en-GB",
            },
        )
        assert resp.status_code == 200
        assert len(resp.content) > 0


def test_docx_rtl_for_arabic_sets_bidi():
    from app.services.export_docx import generate_translation_docx
    data = generate_translation_docx(
        source_text="测试", translated_text="اختبار", risk_annotations=[], language="ar"
    )
    doc = Document(BytesIO(data))
    tr_paras = [p for p in doc.paragraphs if "اختبار" in p.text]
    assert tr_paras, "translation paragraph not found"
    pPr = tr_paras[0]._p.pPr
    assert pPr is not None and pPr.find(qn("w:bidi")) is not None


def test_docx_ltr_for_english_no_bidi():
    from app.services.export_docx import generate_translation_docx
    data = generate_translation_docx(
        source_text="测试", translated_text="test", risk_annotations=[], language="en-GB"
    )
    doc = Document(BytesIO(data))
    tr_paras = [p for p in doc.paragraphs if "test" in p.text]
    assert tr_paras
    pPr = tr_paras[0]._p.pPr
    assert pPr is None or pPr.find(qn("w:bidi")) is None
